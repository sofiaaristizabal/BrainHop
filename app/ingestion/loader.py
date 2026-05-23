"""
loader.py simplemente se encarga de extraer el texto plano de os archivos subidos por el usuario. 
Es el primer paso para ingestar documentos. Devuelve el texto plano de los documentos sin importar su formato. 

Permite:
    - PDF
    - TXT
    - Markdown
    - CSV
    - DOCX
"""

import csv
import io
from pathlib import Path
import pypdf
import docx

def extract_text(file_bytes: bytes, mime_type: str, filename: str):
    """
    Extrae texto plano de un archivo dado sus bytes y su MIME type 

    file_bytes: contenido del archivo como bytes (lo que llega del upload)
    mime_type:  el MIME type validado por el backend (ej: "application/pdf")
    filename:   nombre original del archivo, usado solo para mensajes de error

    retorna string con el contexto completo extraido 
    """

    extractors = {
        "application/pdf": _extract_pdf,
        "text/plain":      _extract_plain_text,
        "text/markdown":   _extract_plain_text,
        "text/csv":        _extract_csv,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    }

    extractor = extractors.get(mime_type)

    if extractor is None: 
        raise ValueError(
            f"File type not supported: '{mime_type}' for file: '{filename}'. "
            f"Allowed file types: {list(extractors.keys())}"
        )
    
    try: 
        raw_text = extractor(file_bytes)
    except Exception as e:
        raise ValueError(
            f"Error extracting the text of the file:  '{filename}': {e}"
        ) from e
    
    #Limpiamos el texto eliminando lineas vacias y espacios al inicio o final del texto
    cleaned = _clean_text(raw_text)

    if not cleaned:
       raise ValueError(
            f"El archivo '{filename}' no contiene texto extraíble. "
            "¿Está el PDF escaneado como imagen? Solo se soportan PDFs con texto seleccionable."
        ) 
    
    return cleaned 

#Ahora vamos a crear los extractores especificos para cada uno de los formatos soportados 

def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extrae texto de un pdf, pero solo funciona su el pdf tiene texto real, no sirve si es escaneado
    """

    pdf_file = io.BytesIO(file_bytes)
    reader = pypdf.PdfReader(pdf_file)

    if len(reader.pages) == 0:
        raise ValueError("El PDF no tiene páginas.")
    
    pages_text = []
    for page_number, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text: #Puede que haya paginas vacias o que solo contengan imagenes
            pages_text.append(page_text)
    
    return "\n\n".join(pages_text)

def _extract_plain_text(file_bytes: bytes) -> str:
    """
    Decodifica archivos de texto plano (.txt) y Markdown (.md).
    """
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")

 
def _extract_csv(file_bytes: bytes) -> str:
    """
    Convierte un CSV a texto legible para el LLM.
    convertimos cada fila a una frase natural:
    "columna1: valor1 | columna2: valor2 | ..."
    """
    text_io = io.StringIO(file_bytes.decode("utf-8"))
    reader = csv.DictReader(text_io)
 
    if reader.fieldnames is None:
        # CSV sin encabezados — tratar como texto plano
        return _extract_plain_text(file_bytes)
 
    rows_as_text = []
    for row in reader:
        # "nombre: Juan | edad: 25 | ciudad: Medellín"
        row_text = " | ".join(
            f"{key}: {value}"
            for key, value in row.items()
            if value  # skip columnas vacías
        )
        if row_text:
            rows_as_text.append(row_text)
 
    return "\n".join(rows_as_text)
 
def _extract_docx(file_bytes: bytes) -> str:
    """
    Extrae texto de archivos Word (.docx) usando python-docx.
    """
    docx_file = io.BytesIO(file_bytes)
    document = docx.Document(docx_file)
 
    parts = []
 
    # Párrafos normales
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
 
    # También podemos extraer tablas, extraemos cada celda como texto
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
 
    return "\n".join(parts)

def _clean_text(text: str) -> str:
    """
    Limpieza básica del texto extraído.
 
    - Elimina líneas en blanco consecutivas (más de 2 seguidas → 1)
    - Elimina espacios al inicio y final del texto total
    - Normaliza saltos de línea de Windows 
    """
    # Normalizar saltos de línea de Windows
    text = text.replace("\r\n", "\n").replace("\r", "\n")
 
    # Colapsar más de 2 líneas en blanco consecutivas a exactamente 1
    lines = text.split("\n")
    cleaned_lines = []
    blank_count = 0
 
    for line in lines:
        if line.strip() == "":
            blank_count += 1
            if blank_count <= 1:  # permitimos máximo 1 línea en blanco seguida
                cleaned_lines.append(line)
        else:
            blank_count = 0
            cleaned_lines.append(line)
 
    return "\n".join(cleaned_lines).strip()